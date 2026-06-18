#!/usr/bin/env python3
"""
Sinh docx theo ĐÚNG mẫu "Khoá Claude Code" (folder D:\\baocao-claude-code):
  - Khối tiêu đề 4 dòng (13pt đậm / 28pt đậm / 13pt / 10pt)
  - "Mục lục" (Heading 1) + TRƯỜNG TOC tự cập nhật của Word
  - Heading 1/2/3 + bullet (List Paragraph)

Dùng: build(out_path, title_lines, body_md)
  title_lines: [(text, size_pt, bold), ...]  (4 dòng)
  body_md: markdown — '# '->H1, '## '->H2, '### '->H3, '- '->bullet, '**x**'->đậm
"""
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _toc(doc):
    """Chèn trường TOC (Word tự dựng mục lục từ Heading 1-3 khi mở/Update Field)."""
    p = doc.add_paragraph()
    run = p.add_run()
    for typ, txt in [("begin", None), (None, ' TOC \\o "1-3" \\h \\z \\u '), ("separate", None),
                     (None, None), ("end", None)]:
        if typ in ("begin", "separate", "end"):
            fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), typ); run._r.append(fld)
        elif txt:
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = txt; run._r.append(it)
    # placeholder text nguoi dung thay khi chua update
    ph = OxmlElement("w:r"); t = OxmlElement("w:t")
    t.text = "(Mở bằng Word → chuột phải vào đây → Update Field để hiện mục lục)"
    ph.append(t); run._r.addprevious(ph) if False else p._p.append(ph)


def _runs(p, text):
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "":
            continue
        r = p.add_run(seg)
        if i % 2 == 1:
            r.bold = True


def _render_md(doc, body_md):
    for raw in body_md.splitlines():
        s = raw.rstrip()
        if not s.strip():
            continue
        if s.startswith("#### "):
            doc.add_heading(s[5:], level=4)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("---"):
            continue
        elif re.match(r"^\s*[-*]\s", s):
            _runs(doc.add_paragraph(style="List Paragraph"), re.sub(r"^\s*[-*]\s", "", s))
        elif re.match(r"^\s*\d+\.\s", s):
            _runs(doc.add_paragraph(style="List Paragraph"), re.sub(r"^\s*\d+\.\s", "", s))
        else:
            _runs(doc.add_paragraph(), s)


def build(out_path, title_lines, body_md, toc=True):
    doc = Document()
    for text, size, bold in title_lines:
        p = doc.add_paragraph()
        r = p.add_run(text); r.bold = bool(bold); r.font.size = Pt(size)
    doc.add_paragraph()
    if toc:
        doc.add_heading("Mục lục", level=1)
        _toc(doc)
    _render_md(doc, body_md)
    doc.save(str(out_path))
    return out_path


# ----- tien ich tao 4 dong tieu de (chia theo DANH MUC that cua khoa) -----
def title_translation(category, n_bai, course):
    return [(category, 13, True), ("Bản dịch tiếng Việt", 28, True),
            (f"Đầy đủ — sát transcript video ({n_bai} bài)", 13, False),
            (f"Khoá học {course} — Bản tiếng Việt", 10, False)]

def title_summary(category, n_bai, course):
    return [(category, 13, True), ("Tóm tắt nội dung", 28, True),
            (f"Tóm tắt chung + {n_bai} bài (theo bản dịch đầy đủ)", 13, False),
            (f"Khoá học {course} — Bản tiếng Việt", 10, False)]

def title_todo(course, n_bai):
    return [(f"Khoá {course}", 13, True), ("Danh sách công việc (To-do)", 28, True),
            (f"Phần A: theo khoá học ({n_bai} bài) • Phần B: tối ưu cho Trường Việt Anh", 13, False),
            (f"Khoá học {course} — Bản tiếng Việt", 10, False)]


if __name__ == "__main__":
    # tao 1 file mau de doi chieu voi reference
    body = """# Tóm tắt chung — Phase 1
Đoạn tóm tắt chung 1.
Đoạn tóm tắt chung 2.

# Tóm tắt từng bài (2 bài)
## Bài 1.1 Tên bài
- Ý thứ nhất, **đậm** một cụm.
- Ý thứ hai.
## Bài 1.2 Tên bài khác
- Ý a.
- Ý b.
"""
    build("E:\\SkoolProject\\BaoCao\\_mau_test.docx", title_summary(1, 2, "Demo"), body)
    print("OK -> _mau_test.docx")
