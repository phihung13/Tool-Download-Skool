#!/usr/bin/env python3
"""
Gop noi dung 1 khoa (mo ta + loi giang/transcript) thanh 1 file de doc / bao cao.

  python export.py --course "freegroup"            # -> courses/freegroup/_TongHop.md
  python export.py --course "freegroup" --docx     # them ban Word (.docx) neu co python-docx
  python export.py                                  # khoa cu SkoolCourse

Khong can thu vien ngoai cho ban Markdown. Ban .docx can python-docx
(pip install python-docx) -> neu thieu se tu bao va bo qua .docx.

Module nay con cho ham gom noi dung (gather_course / lesson_blocks) dung lai o
ai_tools.py (dich / tom tat).
"""
import argparse, re
from pathlib import Path
import config as C

VIDEXT = (".mp4", ".webm", ".mkv", ".mov")


def read_transcript(folder):
    """Loi giang cua 1 bai: uu tien video.txt, roi bat ky *.txt, roi *.srt (bo timestamp)."""
    p = folder / "video.txt"
    if p.exists():
        return p.read_text(encoding="utf-8", errors="replace").strip()
    for q in sorted(folder.glob("*.txt")):
        return q.read_text(encoding="utf-8", errors="replace").strip()
    for q in sorted(folder.glob("*.srt")):
        return srt_to_text(q.read_text(encoding="utf-8", errors="replace"))
    return ""


def srt_to_text(s):
    out = []
    for line in s.splitlines():
        line = line.strip()
        if not line or line.isdigit() or "-->" in line:
            continue
        out.append(line)
    return " ".join(out).strip()


def read_desc(folder):
    p = folder / "description.md"
    return p.read_text(encoding="utf-8", errors="replace").strip() if p.exists() else ""


def _is_lesson(folder):
    """Bai hoc = folder khong co thu muc con (tru 'resources')."""
    subs = [d for d in folder.iterdir() if d.is_dir() and d.name.lower() != "resources"]
    return not subs


def _has_content(folder):
    if read_desc(folder) or read_transcript(folder):
        return True
    return any((folder / ("video" + e)).exists() for e in VIDEXT)


def lesson_blocks(root):
    """Duyet cay khoa theo thu tu folder, tra ve list:
       {depth, title, kind('section'|'lesson'), desc, transcript, path}
       de export.py (xuat) va ai_tools.py (tom tat) dung chung."""
    root = Path(root)
    blocks = []

    def walk(folder, depth):
        kids = sorted([d for d in folder.iterdir() if d.is_dir() and d.name.lower() != "resources"],
                      key=lambda p: p.name)
        for d in kids:
            title = d.name.split(" - ", 1)[-1] if re.match(r"\s*\d+\s*-", d.name) else d.name
            if _is_lesson(d):
                desc, tr = read_desc(d), read_transcript(d)
                if desc or tr or _has_content(d):
                    blocks.append({"depth": depth, "title": title, "kind": "lesson",
                                   "desc": desc, "transcript": tr, "path": d})
            else:
                blocks.append({"depth": depth, "title": title, "kind": "section",
                               "desc": "", "transcript": "", "path": d})
                walk(d, depth + 1)

    walk(root, 1)
    return blocks


def gather_course(root):
    """Tra ve (course_title, blocks)."""
    root = Path(root)
    title = C.COURSE or root.name
    return title, lesson_blocks(root)


# ----------------------------- xuat Markdown -----------------------------
def to_markdown(title, blocks):
    out = [f"# {title}", ""]
    n_lessons = sum(1 for b in blocks if b["kind"] == "lesson")
    out.append(f"*Tổng hợp {n_lessons} bài học — mô tả + lời giảng.*")
    out.append("")
    for b in blocks:
        lvl = min(b["depth"] + 1, 6)
        out.append("#" * lvl + " " + b["title"])
        out.append("")
        if b["kind"] == "lesson":
            if b["desc"]:
                out.append("**Mô tả:**"); out.append(""); out.append(b["desc"]); out.append("")
            if b["transcript"]:
                out.append("**Lời giảng:**"); out.append(""); out.append(b["transcript"]); out.append("")
            if not b["desc"] and not b["transcript"]:
                out.append("*(chưa có mô tả / lời giảng)*"); out.append("")
    return "\n".join(out).rstrip() + "\n"


# ----------------------------- xuat Word (.docx) -------------------------
def to_docx(title, blocks, target):
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        print("[bo qua .docx] Chua cai python-docx -> chay: pip install python-docx")
        return False
    doc = Document()
    doc.add_heading(title, level=0)
    n_lessons = sum(1 for b in blocks if b["kind"] == "lesson")
    doc.add_paragraph(f"Tổng hợp {n_lessons} bài học — mô tả + lời giảng.")
    for b in blocks:
        doc.add_heading(b["title"], level=min(b["depth"], 4))
        if b["kind"] == "lesson":
            if b["desc"]:
                r = doc.add_paragraph(); r.add_run("Mô tả:").bold = True
                for para in b["desc"].split("\n\n"):
                    if para.strip(): doc.add_paragraph(para.strip())
            if b["transcript"]:
                r = doc.add_paragraph(); r.add_run("Lời giảng:").bold = True
                doc.add_paragraph(b["transcript"])
            if not b["desc"] and not b["transcript"]:
                doc.add_paragraph("(chưa có mô tả / lời giảng)")
    doc.save(str(target))
    return True


def run(root=None, docx=False, out=None):
    root = Path(root or C.ROOT)
    title, blocks = gather_course(root)
    if not blocks:
        print("Khong tim thay noi dung (mo ta / transcript) trong khoa nay.\n"); return None
    md_path = Path(out) if out else (root / "_TongHop.md")
    md_path.write_text(to_markdown(title, blocks), encoding="utf-8")
    print(f">> Markdown: {md_path}")
    if docx:
        dx = md_path.with_suffix(".docx")
        if to_docx(title, blocks, dx):
            print(f">> Word: {dx}")
    n = sum(1 for b in blocks if b["kind"] == "lesson")
    print(f"--- EXPORT: {n} bai ---\n")
    return md_path


def main():
    import common as K
    K.setup_console()
    ap = argparse.ArgumentParser(description="Gop noi dung khoa thanh 1 file")
    ap.add_argument("--course", help="Ten khoa duoi courses/. Bo trong = SkoolCourse.")
    ap.add_argument("--docx", action="store_true", help="Xuat them ban Word (.docx).")
    ap.add_argument("--out", help="Duong dan file dau ra (mac dinh <khoa>/_TongHop.md).")
    a = ap.parse_args()
    if a.course: C.set_course(a.course)
    run(root=C.ROOT, docx=a.docx, out=a.out)


if __name__ == "__main__":
    main()
